"""
Export Generation Utilities
Generate PDF and DOCX exports from AI responses

Developed by: DSR
"""

import io
import re
from datetime import datetime
from typing import Optional

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def clean_html_tags(text: str) -> str:
    """Remove HTML tags from text"""
    # Remove HTML tags
    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'<[^>]+>', '', text)
    # Decode HTML entities
    text = text.replace('&lt;', '<')
    text = text.replace('&gt;', '>')
    text = text.replace('&amp;', '&')
    text = text.replace('&quot;', '"')
    return text


def generate_pdf(content: str, store_id: Optional[str] = None) -> bytes:
    """
    Generate PDF from AI response content
    
    Args:
        content: AI response text (may contain HTML)
        store_id: Optional store ID for context
        
    Returns:
        PDF file content as bytes
    """
    if not PDF_AVAILABLE:
        raise ImportError("reportlab not installed")
    
    # Clean HTML tags
    clean_content = clean_html_tags(content)
    
    # Create PDF in memory
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Custom title style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#667eea'),
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    # Custom heading style
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#764ba2'),
        spaceAfter=12,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )
    
    # Custom body style
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=11,
        spaceAfter=12,
        leading=16,
        fontName='Helvetica'
    )
    
    # Add title
    title = Paragraph("V-Mart AI Insights & Recommendations", title_style)
    elements.append(title)
    
    # Add metadata
    metadata_text = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    if store_id:
        metadata_text += f" | Store: {store_id}"
    metadata = Paragraph(metadata_text, styles['Normal'])
    elements.append(metadata)
    elements.append(Spacer(1, 0.3*inch))
    
    # Process content
    lines = clean_content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            elements.append(Spacer(1, 0.1*inch))
            continue
        
        # Check if it's a heading (starts with emoji or **text**)
        if line.startswith('**') and line.endswith('**'):
            heading_text = line.strip('*')
            elements.append(Paragraph(heading_text, heading_style))
        elif any(line.startswith(emoji) for emoji in ['🌤️', '🏪', '📊', '💡', '⚠️', '✅', '📈']):
            elements.append(Paragraph(line, heading_style))
        else:
            elements.append(Paragraph(line, body_style))
    
    # Add footer
    elements.append(Spacer(1, 0.5*inch))
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.gray,
        alignment=TA_CENTER
    )
    footer = Paragraph("Powered by V-Mart AI | Developed by DSR | Inspired by LA", footer_style)
    elements.append(footer)
    
    # Build PDF
    doc.build(elements)
    
    # Get PDF bytes
    pdf_bytes = buffer.getvalue()
    buffer.close()
    
    return pdf_bytes


def generate_docx(content: str, store_id: Optional[str] = None) -> bytes:
    """
    Generate DOCX from AI response content
    
    Args:
        content: AI response text (may contain HTML)
        store_id: Optional store ID for context
        
    Returns:
        DOCX file content as bytes
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed")
    
    # Clean HTML tags
    clean_content = clean_html_tags(content)
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading('V-Mart AI Insights & Recommendations', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    metadata_text = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    if store_id:
        metadata_text += f" | Store: {store_id}"
    metadata_para = doc.add_paragraph(metadata_text)
    metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata_run = metadata_para.runs[0]
    metadata_run.font.size = Pt(10)
    metadata_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # Spacer
    
    # Process content
    lines = clean_content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()  # Empty paragraph for spacing
            continue
        
        # Check if it's a heading
        if line.startswith('**') and line.endswith('**'):
            heading_text = line.strip('*')
            heading = doc.add_heading(heading_text, level=2)
            # Color the heading
            for run in heading.runs:
                run.font.color.rgb = RGBColor(118, 75, 162)  # #764ba2
        elif any(line.startswith(emoji) for emoji in ['🌤️', '🏪', '📊', '💡', '⚠️', '✅', '📈']):
            heading = doc.add_heading(line, level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(102, 126, 234)  # #667eea
        else:
            para = doc.add_paragraph(line)
            para_format = para.paragraph_format
            para_format.space_after = Pt(8)
    
    # Add footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph("Powered by V-Mart AI | Developed by DSR | Inspired by LA")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    
    return docx_bytes
