"""
File Processing Utilities
Extract text and data from various file formats

Developed by: DSR
"""

import io
from typing import Any, Dict, List, Optional

try:
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    import pandas as pd

    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image

    IMAGE_OCR_AVAILABLE = True
except ImportError:
    IMAGE_OCR_AVAILABLE = False


def extract_pdf_text(file_bytes: bytes) -> Dict[str, Any]:
    """
    Extract text from PDF file

    Args:
        file_bytes: PDF file content as bytes

    Returns:
        Dictionary with extracted text and metadata
    """
    if not PDF_AVAILABLE:
        return {"success": False, "error": "PyPDF2 not installed"}

    try:
        pdf_file = io.BytesIO(file_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        text_content = []
        for page_num, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            if text.strip():
                text_content.append(f"[Page {page_num + 1}]\n{text}")

        full_text = "\n\n".join(text_content)

        return {
            "success": True,
            "text": full_text,
            "page_count": len(pdf_reader.pages),
            "char_count": len(full_text),
            "method": "PyPDF2",
        }
    except Exception as e:
        return {"success": False, "error": f"PDF extraction failed: {str(e)}"}


def extract_docx_text(file_bytes: bytes) -> Dict[str, Any]:
    """
    Extract text from DOCX file

    Args:
        file_bytes: DOCX file content as bytes

    Returns:
        Dictionary with extracted text and metadata
    """
    if not DOCX_AVAILABLE:
        return {"success": False, "error": "python-docx not installed"}

    try:
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)

        # Extract paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Extract tables
        tables_text = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(" | ".join(row_data))
            if table_data:
                tables_text.append("\n".join(table_data))

        full_text = "\n\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n[TABLES]\n" + "\n\n".join(tables_text)

        return {
            "success": True,
            "text": full_text,
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
            "char_count": len(full_text),
            "method": "python-docx",
        }
    except Exception as e:
        return {"success": False, "error": f"DOCX extraction failed: {str(e)}"}


def extract_excel_data(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    Extract data from Excel/CSV file (supports multiple sheets)

    Args:
        file_bytes: Excel/CSV file content as bytes
        filename: Original filename to determine format

    Returns:
        Dictionary with extracted data and metadata
    """
    if not EXCEL_AVAILABLE:
        return {"success": False, "error": "openpyxl/pandas not installed"}

    try:
        is_csv = filename.lower().endswith(".csv")

        if is_csv:
            # Parse CSV - preserve string values, don't auto-convert
            csv_file = io.StringIO(file_bytes.decode("utf-8"))
            df = pd.read_csv(csv_file, dtype=str, keep_default_na=False)
            sheets_data = {"Sheet1": df}
        else:
            # Parse Excel - READ ALL SHEETS
            excel_file = io.BytesIO(file_bytes)
            excel_obj = pd.ExcelFile(excel_file)
            sheets_data = {}

            # Read all sheets
            for sheet_name in excel_obj.sheet_names:
                sheets_data[sheet_name] = pd.read_excel(
                    excel_obj, sheet_name=sheet_name
                )

        # Build comprehensive text from all sheets
        text_parts = []
        all_summaries = {}

        for sheet_name, df in sheets_data.items():
            # Get summary for this sheet
            summary = {
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": df.columns.tolist(),
                "preview": df.head(10).to_dict(orient="records"),
            }
            all_summaries[sheet_name] = summary

            # Add sheet header
            if len(sheets_data) > 1:
                text_parts.append(f"\n{'=' * 60}")
                text_parts.append(f"SHEET: {sheet_name}")
                text_parts.append(f"{'=' * 60}")

            # Convert to readable text - use full data representation
            text_parts.append(
                f"Data Summary: {summary['rows']} rows × {summary['columns']} columns"
            )
            text_parts.append(f"Columns: {', '.join(summary['column_names'])}")

            # Add RAW CSV-like representation (all rows, preserving original values)
            text_parts.append("\n📊 COMPLETE DATA (All Rows):")
            text_parts.append("-" * 80)

            # For CSV files, include raw CSV format for better AI understanding
            if is_csv:
                # Add header
                text_parts.append(",".join(summary["column_names"]))
                # Add all data rows (not just preview)
                for _, row in df.iterrows():
                    text_parts.append(",".join(str(val) for val in row.values))
                text_parts.append("-" * 80)

            # Also add table format for better readability
            text_parts.append("\n📋 TABLE FORMAT (First 20 rows):")
            text_parts.append(df.head(20).to_string(index=False))

            # Add basic statistics for numeric columns (if any exist after conversion)
            try:
                numeric_df = df.apply(pd.to_numeric, errors="ignore")
                numeric_cols = numeric_df.select_dtypes(
                    include=["number"]
                ).columns.tolist()
                if numeric_cols:
                    text_parts.append("\n📈 NUMERIC COLUMN STATISTICS:")
                    text_parts.append(numeric_df[numeric_cols].describe().to_string())
            except Exception:
                pass  # Skip stats if conversion fails

        full_text = "\n".join(text_parts)

        return {
            "success": True,
            "text": full_text,
            "sheets": list(sheets_data.keys()),
            "sheet_count": len(sheets_data),
            "data_summary": all_summaries,
            "char_count": len(full_text),
            "method": "pandas_multi_sheet",
        }
    except Exception as e:
        return {"success": False, "error": f"Excel/CSV extraction failed: {str(e)}"}


def extract_text_file(file_bytes: bytes) -> Dict[str, Any]:
    """
    Extract text from plain text file

    Args:
        file_bytes: Text file content as bytes

    Returns:
        Dictionary with extracted text and metadata
    """
    try:
        # Try UTF-8 first, then fallback to latin-1
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1")

        return {
            "success": True,
            "text": text,
            "char_count": len(text),
            "line_count": text.count("\n") + 1,
            "method": "text",
        }
    except Exception as e:
        return {"success": False, "error": f"Text extraction failed: {str(e)}"}


def extract_image_text(file_bytes: bytes) -> Dict[str, Any]:
    """
    Extract text from image using OCR (if available)

    Args:
        file_bytes: Image file content as bytes

    Returns:
        Dictionary with extracted text and metadata
    """
    if not IMAGE_OCR_AVAILABLE:
        return {
            "success": True,
            "text": "[Image file attached - OCR not available for text extraction]",
            "note": "Image can be analyzed by Gemini Vision API",
            "method": "placeholder",
        }

    try:
        image = Image.open(io.BytesIO(file_bytes))

        # Try OCR extraction
        try:
            text = pytesseract.image_to_string(image)
            if text.strip():
                return {
                    "success": True,
                    "text": text,
                    "char_count": len(text),
                    "image_size": image.size,
                    "method": "OCR (pytesseract)",
                }
        except Exception:
            pass

        # If OCR fails or no text, return image info
        return {
            "success": True,
            "text": f"[Image file: {image.size[0]}x{image.size[1]} pixels, format: {image.format}]",
            "note": "Image can be analyzed by Gemini Vision API",
            "image_size": image.size,
            "image_format": image.format,
            "method": "image_metadata",
        }
    except Exception as e:
        return {"success": False, "error": f"Image processing failed: {str(e)}"}


def process_uploaded_file(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    Process uploaded file and extract content based on file type

    Args:
        file_bytes: File content as bytes
        filename: Original filename

    Returns:
        Dictionary with extracted content and metadata
    """
    ext = filename.lower().split(".")[-1]

    processors = {
        "pdf": extract_pdf_text,
        "docx": extract_docx_text,
        "doc": extract_docx_text,
        "xlsx": lambda b: extract_excel_data(b, filename),
        "xls": lambda b: extract_excel_data(b, filename),
        "csv": lambda b: extract_excel_data(b, filename),
        "txt": extract_text_file,
        "jpg": extract_image_text,
        "jpeg": extract_image_text,
        "png": extract_image_text,
        "gif": extract_image_text,
    }

    processor = processors.get(ext)

    if not processor:
        return {"success": False, "error": f"Unsupported file type: .{ext}"}

    result = processor(file_bytes)
    result["filename"] = filename
    result["file_type"] = ext

    return result
